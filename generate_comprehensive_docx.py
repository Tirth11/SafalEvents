# -*- coding: utf-8 -*-
"""Generate Comprehensive End-to-End System Requirements Word document for SafalEvents."""
import datetime
import sys
import os

# Set cwd or insert path to ensure imports work
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Input/Output paths
OUT = "/Users/tirth/Documents/SafalEvents_Mockup_Web/SafalEvents_Comprehensive_Requirements.docx"
TODAY = datetime.date(2026, 6, 27).strftime("%d %B %Y")

# Branding Colors
BRAND = RGBColor(0x1E, 0x3A, 0x8A)      # Deep blue
INK   = RGBColor(0x11, 0x18, 0x27)      # Very dark gray
MUTED = RGBColor(0x4B, 0x55, 0x63)      # Dark gray
LABELBG = "F3F4F6"                      # Light gray
HEADBG  = "1E3A8A"                      # Deep blue header background

# Import single source of truth data
from build_user_stories import EPICS
from build_sprint_plan import FEATURES

# Build story-to-feature mapping
story_to_feature = {}
feature_count_map = {}
for epic_id, feat_list in FEATURES.items():
    feature_count_map[epic_id] = len(feat_list)
    for feat_id, feat_name, story_ids in feat_list:
        for sid in story_ids:
            story_to_feature[sid] = (feat_id, feat_name)

# ----------------------------------------------------------------------------
# DOCX BUILD HELPERS
# ----------------------------------------------------------------------------
def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hexcolor)
    tcPr.append(shd)

def set_cell_text(cell, text, bold=False, color=None, size=None, white=False):
    cell.text = ''
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    run.bold = bold
    if size:
        run.font.size = Pt(size)
    if white:
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    elif color:
        run.font.color.rgb = color
    else:
        run.font.color.rgb = INK
    return p

def add_bullets(cell_or_doc, items):
    if hasattr(cell_or_doc, 'text'):
        # It's a cell
        cell_or_doc.text = ''
        for i, it in enumerate(items):
            p = cell_or_doc.paragraphs[0] if i == 0 else cell_or_doc.add_paragraph()
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.line_spacing = 1.15
            p.style = doc.styles['List Bullet']
            p.add_run(it)
    else:
        # It's a document
        for it in items:
            p = cell_or_doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.line_spacing = 1.15
            p.add_run(it)

def add_paragraph_styled(doc, text="", bold=False, color=None, size=None, align=None, space_after=6, space_before=0, style=None):
    if style:
        p = doc.add_paragraph(style=style)
    else:
        p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing = 1.15
    if align:
        p.alignment = align
    if text:
        run = p.add_run(text)
        run.bold = bold
        if size:
            run.font.size = Pt(size)
        if color:
            run.font.color.rgb = color
    return p

def add_toc(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    
    run._r.append(fldChar)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)

# ----------------------------------------------------------------------------
# DYNAMIC TEST SCENARIOS GENERATOR
# ----------------------------------------------------------------------------
def get_story_scenarios(story):
    role = story['role']
    title = story['title']
    
    positives = []
    negatives = []
    boundaries = []
    permissions = []
    validations = []
    
    # Defaults
    positives.append(f"Verify that a user as '{role}' can successfully complete '{title}' and retrieve all success confirmation states.")
    negatives.append(f"Verify that unauthorized attempts to access or execute '{title}' fail with an appropriate error state.")
    boundaries.append(f"Verify that the system handles resource limit boundaries for '{title}' without freezing or double processing.")
    permissions.append(f"Verify that only verified '{role}' roles have access to execute '{title}' (Default-Deny policy).")
    validations.append(f"Verify input bounds and display validation alerts for empty, null, or incorrect formats during '{title}'.")
    
    # Specific overrides to make it super rich!
    if "OTP" in title or "verification" in title.lower():
        positives = [
            "Verify that a 6-digit OTP code is sent to the registered email/phone on triggering verification.",
            "Verify that entering the correct active OTP successfully completes the verification process."
        ]
        negatives = [
            "Verify that entering an incorrect OTP code 5 times locks/invalidates the verification session.",
            "Verify that entering an expired OTP (older than 10 minutes) is rejected with a clear message and a resend button."
        ]
        boundaries = [
            "Verify behavior at exactly the 5th failed attempt limit (session should lock exactly on the 5th attempt).",
            "Verify behavior exactly at 10 minutes and 0 seconds of OTP expiry time."
        ]
        permissions = [
            "Verify that verification status is updated only for the session owner.",
            "Verify that an unauthenticated API request cannot bypass the OTP check."
        ]
        validations = [
            "Verify that non-numeric characters in the OTP input field are rejected on submission.",
            "Verify that empty OTP submissions display a validation warning and do not hit the server."
        ]
    elif "sign up" in title.lower() or "register" in title.lower():
        positives = [
            "Verify that a new host/guest can register successfully with valid details (name, email, phone, password).",
            "Verify that the host profile starts in the correct default state ('ACTIVE' for individuals, 'PENDING_ADMIN_APPROVAL' for organizations)."
        ]
        negatives = [
            "Verify that registration fails and shows an error if the email address is already in use.",
            "Verify that organization signup fails if the website URL is invalid or the website is down."
        ]
        boundaries = [
            "Verify document upload size limits (e.g., uploading exactly at the 5MB file size limit).",
            "Verify character limits for name (e.g., 50 characters max) and organization website URL."
        ]
        permissions = [
            "Verify that pending organization hosts cannot publish events until Superadmin approval.",
            "Verify that guests cannot access host-only onboarding forms."
        ]
        validations = [
            "Verify validation on email field (must contain '@' and domain).",
            "Verify that passwords must meet the minimum strength requirements (e.g., 8+ chars, numbers, symbols)."
        ]
    elif "capacity" in title.lower() or "limit" in title.lower() or "waitlist" in title.lower() or "deadline" in title.lower():
        positives = [
            "Verify that when capacity is available, a new guest RSVP is marked as 'going'.",
            "Verify that when capacity is full, new RSVPs are automatically routed to the waitlist."
        ]
        negatives = [
            "Verify that guests cannot bypass the waitlist status when the event is at full capacity.",
            "Verify that reducing event capacity below the current confirmed attendee count is blocked on the host dashboard."
        ]
        boundaries = [
            "Verify that the guest exactly at capacity is registered as 'going', and the next guest (capacity + 1) is waitlisted.",
            "Verify behavior when registering exactly at the RSVP deadline timestamp (down to the second)."
        ]
        permissions = [
            "Verify that only the host/authorized staff can promote waitlisted guests manually.",
            "Verify that waitlisted guests cannot check in at the gate."
        ]
        validations = [
            "Verify that non-integer values for guest count are rejected.",
            "Verify that waitlist status changes are validated and logged to the audit trail."
        ]
    elif "approve" in title.lower() or "decline" in title.lower() or "reject" in title.lower():
        positives = [
            "Verify that approving a pending RSVP changes their status to 'going' and sends a confirmation email.",
            "Verify that rejecting a pending RSVP changes their status to 'declined/rejected' and displays the rejection reason."
        ]
        negatives = [
            "Verify that an RSVP cannot be approved if the event is already at capacity (unless manually overridden).",
            "Verify that a rejected RSVP cannot be checked in."
        ]
        boundaries = [
            "Verify bulk approval of RSVPs matching the exact remaining event capacity.",
            "Verify character limit of the rejection reason text field."
        ]
        permissions = [
            "Verify that only users with the 'guests_approve' permission can approve/reject RSVPs.",
            "Verify that staff with read-only roles cannot see approval buttons."
        ]
        validations = [
            "Verify validation on rejection reason (e.g., must not contain malicious scripts).",
            "Verify that approval status changes update the audit trail accurately."
        ]
    elif "payment" in title.lower() or "paid" in title.lower() or "ticket" in title.lower():
        positives = [
            "Verify that enabling paid tickets requires a price to be set.",
            "Verify that completing a payment processes the transaction and registers the guest."
        ]
        negatives = [
            "Verify that a payment with invalid card details is rejected and the RSVP is not confirmed.",
            "Verify that an RSVP is not recorded as 'going' if payment fails."
        ]
        boundaries = [
            "Verify processing a payment with ticket price set to $0.01 (minimum boundary).",
            "Verify calculations when RSVPing with multiple additional guests (total = price * count)."
        ]
        permissions = [
            "Verify that only the host (event owner) can configure payment details and bank settings.",
            "Verify that payment transactions are secure and unmodifiable by guests."
        ]
        validations = [
            "Verify credit card number, expiry date (MM/YY), and CVC format validation.",
            "Verify that negative values are rejected in the ticket price field."
        ]
    elif "scan" in title.lower() or "check-in" in title.lower() or "checkin" in title.lower():
        positives = [
            "Verify that scanning a valid guest QR code successfully checks them in and changes status to 'Arrived'.",
            "Verify that check-in details (scanner name, timestamp) are recorded in the event history."
        ]
        negatives = [
            "Verify that scanning an already checked-in QR code displays a warning 'Already scanned'.",
            "Verify that scanning an invalid or wrong event QR code displays 'Not valid for this event'."
        ]
        boundaries = [
            "Verify scanning at the exact event start time.",
            "Verify checking in a guest whose party size is greater than 1 (multiple attendees under one booking)."
        ]
        permissions = [
            "Verify that only users with 'checkin' permission (e.g., Host or QR Scanner staff) can perform check-in.",
            "Verify that guests cannot access the scanner module."
        ]
        validations = [
            "Verify check-in timestamp format is correct in the logs.",
            "Verify QR code decrypts to a valid booking ID."
        ]
    elif "staff" in title.lower() or "role" in title.lower() or "permission" in title.lower():
        positives = [
            "Verify that the host can invite a staff member via email and assign them a specific role.",
            "Verify that a staff member logging in with a valid Invite ID sees only their permitted screens."
        ]
        negatives = [
            "Verify that a staff member cannot access non-permitted screens via direct URL manipulation.",
            "Verify that an expired or revoked Invite ID is rejected at login."
        ]
        boundaries = [
            "Verify inviting multiple staff members with the same role.",
            "Verify behaviour when a role is configured with zero permissions (default-deny)."
        ]
        permissions = [
            "Verify that only the Host (or staff with 'staff_manage' permission) can invite/revoke staff.",
            "Verify that staff cannot modify or elevate their own roles."
        ]
        validations = [
            "Verify email format validation for staff invitation.",
            "Verify Invite ID format and validation logs."
        ]
        
    return positives, negatives, boundaries, permissions, validations


# ----------------------------------------------------------------------------
# DOCUMENT GENERATION
# ----------------------------------------------------------------------------
doc = Document()

# Base Font Setup
normal = doc.styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(11)
normal.font.color.rgb = INK

# Headings Font Setup
for h, sz in [('Heading 1', 18), ('Heading 2', 14), ('Heading 3', 12), ('Title', 32)]:
    st = doc.styles[h]
    st.font.name = 'Calibri'
    st.font.bold = True
    if h != 'Title':
        st.font.color.rgb = BRAND

# ----------------- 1. TITLE & COVER PAGE -----------------
add_paragraph_styled(doc, "\n\n\n\n", space_after=12)
add_paragraph_styled(doc, "SafalEvents", bold=True, color=BRAND, size=40, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)
add_paragraph_styled(doc, "Comprehensive End-to-End System Requirements", bold=True, size=20, color=INK, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)
add_paragraph_styled(doc, "Hosts, Guests, Superadmin, and Staff Functionality Document", size=12, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=36)

add_paragraph_styled(doc, "\n", space_after=12)

# Document Information Table
info_tbl = doc.add_table(rows=10, cols=2)
info_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
info_tbl.style = 'Table Grid'
info_tbl.autofit = False
info_tbl.columns[0].width = Inches(2.0)
info_tbl.columns[1].width = Inches(4.5)

doc_info = [
    ("Product Name", "SafalEvents — Event hosting, RSVP & ticketing platform"),
    ("Module Name", "Hosts, Guests, Superadmin, Staff Comprehensive Requirements"),
    ("Document Version", "2.0.0"),
    ("Author", "Project Lead Analyst"),
    ("Last Updated", TODAY),
    ("Document Status", "Approved"),
    ("Target Platforms", "Web / Android / iOS (React Native / Expo build)"),
    ("Related Modules", "Core Authentication, RSVP approval engine, Gate Scanner & Staff RBAC"),
    ("Web Mockup URL", "https://safal-events.vercel.app"),
    ("Mobile Mockup URL", "https://safalevents.netlify.app")
]

for idx, (k, v) in enumerate(doc_info):
    info_tbl.rows[idx].height_rule = WD_ROW_HEIGHT_RULE.AUTO
    set_cell_text(info_tbl.rows[idx].cells[0], k, bold=True)
    shade(info_tbl.rows[idx].cells[0], LABELBG)
    set_cell_text(info_tbl.rows[idx].cells[1], v)

doc.add_page_break()

# ----------------- TABLE OF CONTENTS -----------------
doc.add_heading("Table of Contents", level=1)
add_paragraph_styled(doc, "Note: If the Table of Contents does not update automatically, right-click and select 'Update Field'.", size=9, color=MUTED, space_after=12)
add_toc(doc)
doc.add_page_break()

# ----------------- 2. REVISION HISTORY -----------------
doc.add_heading("1. Revision History", level=1)
rev_tbl = doc.add_table(rows=3, cols=4)
rev_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
rev_tbl.style = 'Table Grid'
rev_tbl.autofit = False
rev_tbl.columns[0].width = Inches(1.0)
rev_tbl.columns[1].width = Inches(1.2)
rev_tbl.columns[2].width = Inches(1.5)
rev_tbl.columns[3].width = Inches(3.0)

rev_hdrs = ["Version", "Date", "Author", "Change Description"]
for col_idx, hdr in enumerate(rev_hdrs):
    set_cell_text(rev_tbl.rows[0].cells[col_idx], hdr, bold=True, white=True)
    shade(rev_tbl.rows[0].cells[col_idx], HEADBG)

rev_history = [
    ("1.0.0", "18 June 2026", "Project Lead", "Initial baseline user stories & sprint mapping."),
    ("2.0.0", "27 June 2026", "Lead QA & Analyst", "Expanded document to include exhaustive UAT scenarios, boundary conditions, validation schemas, and web/mobile mockup references.")
]

for row_idx, data in enumerate(rev_history):
    for col_idx, text in enumerate(data):
        set_cell_text(rev_tbl.rows[row_idx + 1].cells[col_idx], text)

add_paragraph_styled(doc, "\n")

# ----------------- 3. OVERVIEW -----------------
doc.add_heading("2. Overview", level=1)
doc.add_heading("2.1 Business Overview", level=2)
doc.add_paragraph(
    "SafalEvents is an enterprise-grade event hosting, RSVP, and ticketing platform. "
    "Designed to service high-volume and high-trust events, it bridges the gaps between hosts (organizers), "
    "guests (attendees), Superadmins (platform operators), and temporary staff. "
    "The application relies on a unified backend powering two responsive client interfaces: a Web-based platform "
    "geared toward desktop host operations and event page discoveries, and a React Native mobile frame specifically "
    "optimized for check-in staff at the venue gate, guest access, and real-time host notifications."
)

doc.add_heading("2.2 Problem Statement", level=2)
doc.add_paragraph(
    "Traditional event platforms suffer from ticket fraud, identity spoofing, and lack of granular staff management. "
    "Hosts are forced to use multiple, disconnected tools to handle RSVP approvals, gate scanner check-ins, custom forms, "
    "and payouts. Furthermore, temporary staff are often granted over-privileged access to event dashboard databases. "
    "SafalEvents implements a default-deny, role-based permission system scoped per event, with integrated OTP verification "
    "and waitlisting, eliminating standard ticketing leaks and organizational vulnerabilities."
)

doc.add_heading("2.3 Business Goals", level=2)
doc.add_paragraph(
    "1. High Trust: Ensure all event RSVPs represent verified email/phone endpoints via automated OTP validation.\n"
    "2. Security & RBAC: Shield critical customer information by providing scoped sub-dashboards to staff helpers.\n"
    "3. Scalability & Efficiency: Automate approval queues and waitlists FIFO-style to maintain full venue capacity."
)

doc.add_heading("2.4 Objectives", level=2)
doc.add_paragraph(
    "• Deliver cross-platform parity across Web (https://safal-events.vercel.app) and mobile frame (https://safalevents.netlify.app).\n"
    "• Establish immutable audit trails tracking guest check-ins, manual invitations, and role configurations.\n"
    "• Support secure in-app messaging, interactive comments moderation, and real-time voting on event schedules."
)


# ----------------- 4. SCOPE -----------------
doc.add_heading("3. Scope", level=1)
doc.add_heading("3.1 In Scope", level=2)
scope_items = [
    "Unified login scoping (Guest, Host, Admin, and Staff via Invite-ID).",
    "Event creation, editing, cloning, draft saving, and deleting.",
    "Granular RSVP rules: capacity caps, approval flags, age restrictions, and self-service cancellations.",
    "OTP-based authentication engine and email/SMS confirmation notifications.",
    "Integrated paid ticketing gating and bank payout registrations.",
    "Host activity feed inbox and notification outbox history logs.",
    "Superadmin Console with user oversight, host approvals, suspensions, and platform color configurations.",
    "Dynamic event-based staff invitation and custom role creator (RBAC).",
    "QR Pass generation and gate scanning functionality with an offline parity check.",
    "Event page engagement: comments feed, reactions, pin moderation, and real-time polls."
]
doc.add_paragraph("The system encapsulates the following core functional areas:")
for item in scope_items:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(item)

doc.add_heading("3.2 Out of Scope", level=2)
out_scope_items = [
    "Multi-currency conversion (the system resolves transactions strictly in USD/fixed currency).",
    "Physical ticket printing and post-delivery (resolved via digital QR Pass check-ins).",
    "Advanced CRM campaign automation (messaging is limited to event updates, reminders, and manual broadcasts)."
]
for item in out_scope_items:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(item)

doc.add_page_break()

# ----------------- 5. TARGET AUDIENCE & USER ROLES -----------------
doc.add_heading("4. Target Audience & User Roles", level=1)
doc.add_paragraph(
    "The application defines four distinct user personas, each operating within a strict security scope:"
)

roles_data = [
    ("Host (Organizer)", "Creates and manages events. Sets up ticketing, approval parameters, custom forms, and invitations. Invites and assigns roles to Staff. Monitored for platform policy compliance by Superadmins."),
    ("Guest (Attendee)", "Browses public events, signs up, verifies identity via OTP, makes ticket payments, receives digital passes, cancels or edits RSVPs, messages hosts, comments, and votes on polls."),
    ("Staff (Helper)", "Enters an event space via a specific Invite ID. Scoped strictly to permitted tasks assigned by the Host (e.g., QR check-ins, guest viewing, or messaging) without accessing settings."),
    ("Superadmin (Platform Owner)", "Accesses the Admin Console. Oversees users, reviews organization host credentials, approves or suspends accounts, reviews system templates, and audits security logs.")
]

role_tbl = doc.add_table(rows=5, cols=2)
role_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
role_tbl.style = 'Table Grid'
role_tbl.autofit = False
role_tbl.columns[0].width = Inches(1.8)
role_tbl.columns[1].width = Inches(4.7)

set_cell_text(role_tbl.rows[0].cells[0], "User Role / Persona", bold=True, white=True)
shade(role_tbl.rows[0].cells[0], HEADBG)
set_cell_text(role_tbl.rows[0].cells[1], "Description and Operational Boundaries", bold=True, white=True)
shade(role_tbl.rows[0].cells[1], HEADBG)

for idx, (role_name, desc) in enumerate(roles_data):
    set_cell_text(role_tbl.rows[idx+1].cells[0], role_name, bold=True)
    shade(role_tbl.rows[idx+1].cells[0], LABELBG)
    set_cell_text(role_tbl.rows[idx+1].cells[1], desc)

doc.add_page_break()

# ----------------- 6. EPIC SUMMARY -----------------
doc.add_heading("5. Epic Summary", level=1)
epic_tbl = doc.add_table(rows=len(EPICS) + 1, cols=5)
epic_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
epic_tbl.style = 'Table Grid'
epic_tbl.autofit = False
epic_tbl.columns[0].width = Inches(1.0)
epic_tbl.columns[1].width = Inches(1.5)
epic_tbl.columns[2].width = Inches(2.6)
epic_tbl.columns[3].width = Inches(0.7)
epic_tbl.columns[4].width = Inches(0.7)

headers = ["Epic ID", "Epic Name", "Description", "Total Features", "Total Stories"]
for idx, h in enumerate(headers):
    set_cell_text(epic_tbl.rows[0].cells[idx], h, bold=True, white=True)
    shade(epic_tbl.rows[0].cells[idx], HEADBG)

for row_idx, ep in enumerate(EPICS):
    cells = epic_tbl.rows[row_idx + 1].cells
    set_cell_text(cells[0], "EPIC-" + ep['id'], bold=True)
    set_cell_text(cells[1], ep['name'])
    set_cell_text(cells[2], ep['desc'])
    set_cell_text(cells[3], str(feature_count_map.get(ep['id'], 0)))
    set_cell_text(cells[4], str(len(ep['stories'])))

add_paragraph_styled(doc, "\n")

# ----------------- 7. FEATURE SUMMARY -----------------
doc.add_heading("6. Feature Summary", level=1)
feat_rows = []
for epic_id, f_list in FEATURES.items():
    for f_id, f_name, s_ids in f_list:
        feat_rows.append((f_id, f_name, "EPIC-" + epic_id, len(s_ids)))

feat_tbl = doc.add_table(rows=len(feat_rows) + 1, cols=4)
feat_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
feat_tbl.style = 'Table Grid'
feat_tbl.autofit = False
feat_tbl.columns[0].width = Inches(1.5)
feat_tbl.columns[1].width = Inches(3.0)
feat_tbl.columns[2].width = Inches(1.2)
feat_tbl.columns[3].width = Inches(0.8)

f_headers = ["Feature ID", "Feature Name", "Epic Mapping", "Stories"]
for idx, h in enumerate(f_headers):
    set_cell_text(feat_tbl.rows[0].cells[idx], h, bold=True, white=True)
    shade(feat_tbl.rows[0].cells[idx], HEADBG)

for row_idx, row_data in enumerate(feat_rows):
    cells = feat_tbl.rows[row_idx + 1].cells
    set_cell_text(cells[0], row_data[0], bold=True)
    set_cell_text(cells[1], row_data[1])
    set_cell_text(cells[2], row_data[2])
    set_cell_text(cells[3], str(row_data[3]))

doc.add_page_break()

# ----------------- 8. USER STORY SUMMARY -----------------
doc.add_heading("7. User Story Summary", level=1)
total_stories_count = 0
story_rows = []
for ep in EPICS:
    for s in ep['stories']:
        total_stories_count += 1
        # Map epic to platform
        platform = "Web"
        if ep['id'] == "MOBILE":
            platform = "Android / iOS"
        elif ep['id'] in ["AUTH", "ACCESS", "RSVP", "GUEST", "GM", "CHECKIN"]:
            platform = "Web / App"
        
        story_rows.append((s['id'], s['title'], s['pri'], platform, "Implemented"))

story_tbl = doc.add_table(rows=len(story_rows) + 1, cols=5)
story_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
story_tbl.style = 'Table Grid'
story_tbl.autofit = False
story_tbl.columns[0].width = Inches(1.4)
story_tbl.columns[1].width = Inches(2.6)
story_tbl.columns[2].width = Inches(0.8)
story_tbl.columns[3].width = Inches(1.0)
story_tbl.columns[4].width = Inches(0.9)

s_headers = ["User Story ID", "User Story Name", "Priority", "Platform", "Status"]
for idx, h in enumerate(s_headers):
    set_cell_text(story_tbl.rows[0].cells[idx], h, bold=True, white=True)
    shade(story_tbl.rows[0].cells[idx], HEADBG)

for row_idx, row_data in enumerate(story_rows):
    cells = story_tbl.rows[row_idx + 1].cells
    set_cell_text(cells[0], row_data[0], bold=True)
    set_cell_text(cells[1], row_data[1])
    set_cell_text(cells[2], row_data[2])
    set_cell_text(cells[3], row_data[3])
    set_cell_text(cells[4], row_data[4])

doc.add_page_break()

# ----------------- 9. DETAILED USER STORIES & UAT -----------------
doc.add_heading("8. Detailed User Stories", level=1)

story_counter = 0
for ep in EPICS:
    doc.add_heading(f"EPIC-{ep['id']} — {ep['name']}", level=2)
    doc.add_paragraph(ep['desc'])
    
    for s in ep['stories']:
        story_counter += 1
        
        # User Story Heading
        doc.add_heading(f"{s['id']} — {s['title']}", level=3)
        
        # Get mapped feature info
        mapped_feat = story_to_feature.get(s['id'], ("—", "—"))
        
        # Derive platforms
        platform = "Web"
        if ep['id'] == "MOBILE":
            platform = "Android / iOS (React Native)"
        elif ep['id'] in ["AUTH", "ACCESS", "RSVP", "GUEST", "GM", "CHECKIN", "MSG"]:
            platform = "Web / Android / iOS"
            
        # Persona
        persona = s['role'].capitalize()
        
        # Story Points based on priority
        sp = "5" if s['pri'] == "High" else "3" if s['pri'] == "Medium" else "2"
        
        # 1. User Story Information Table
        tbl = doc.add_table(rows=10, cols=2)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        tbl.style = 'Table Grid'
        tbl.autofit = False
        tbl.columns[0].width = Inches(2.2)
        tbl.columns[1].width = Inches(4.3)
        
        def add_info_row(row_idx, label, val):
            c0 = tbl.rows[row_idx].cells[0]
            c1 = tbl.rows[row_idx].cells[1]
            set_cell_text(c0, label, bold=True)
            shade(c0, LABELBG)
            set_cell_text(c1, val)
            
        add_info_row(0, "User Story ID", s['id'])
        add_info_row(1, "User Story Name", s['title'])
        add_info_row(2, "Epic ID", "EPIC-" + ep['id'])
        add_info_row(3, "Feature ID", mapped_feat[0])
        add_info_row(4, "Priority", s['pri'])
        add_info_row(5, "Story Points", sp)
        add_info_row(6, "User Persona / User Role", persona)
        add_info_row(7, "Platform", platform)
        add_info_row(8, "Dependencies", s['deps'])
        add_info_row(9, "Related User Stories", "—" if s['deps'] == "—" else s['deps'])
        
        doc.add_paragraph("") # Space
        
        # 2. User Story Description
        add_paragraph_styled(doc, "User Story Description", bold=True, size=11, space_after=3)
        doc.add_paragraph(f"As a {s['role']}, I want to {s['want']}, so that {s['benefit']}.")
        
        # 3. Acceptance Criteria
        add_paragraph_styled(doc, "User Acceptance Criteria (UAT)", bold=True, size=11, space_after=3)
        formatted_ac = [f"AC-{i+1}: {ac}" for i, ac in enumerate(s['ac'])]
        add_bullets(doc, formatted_ac)
        
        # 4. Test Scenarios
        add_paragraph_styled(doc, "Test Scenarios", bold=True, size=11, space_after=3)
        
        # Generate scenarios
        positives, negatives, boundaries, permissions, validations = get_story_scenarios(s)
        
        # Render lists for each scenario type
        doc.add_paragraph().add_run("Positive Test Cases:").bold = True
        add_bullets(doc, positives)
        
        doc.add_paragraph().add_run("Negative Test Cases:").bold = True
        add_bullets(doc, negatives)
        
        doc.add_paragraph().add_run("Boundary Test Cases:").bold = True
        add_bullets(doc, boundaries)
        
        doc.add_paragraph().add_run("Permission Test Cases:").bold = True
        add_bullets(doc, permissions)
        
        doc.add_paragraph().add_run("Validation Test Cases:").bold = True
        add_bullets(doc, validations)
        
        doc.add_paragraph("") # Separator space

# ---------- FOOTER & PAGE NUMBERS ----------
def add_page_number(footer_par):
    run = footer_par.add_run()
    fldStart = OxmlElement('w:fldChar')
    fldStart.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = 'PAGE'
    fldEnd = OxmlElement('w:fldChar')
    fldEnd.set(qn('w:fldCharType'), 'end')
    run._r.append(fldStart)
    run._r.append(instr)
    run._r.append(fldEnd)

sec = doc.sections[0]
foot = sec.footer.paragraphs[0]
foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = foot.add_run('SafalEvents — Comprehensive End-to-End System Requirements Document  ·  Page ')
fr.font.size = Pt(8)
fr.font.color.rgb = MUTED
add_page_number(foot)

# Save Document
doc.save(OUT)
print("Finished! Total user stories built:", total_stories_count)
print("Saved to:", OUT)
