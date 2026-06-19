# -*- coding: utf-8 -*-
"""Generate End-to-End User Stories Word document for SafalEvents."""
import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = r"E:\Safalvir\SafalEvents\SafalEvents_User_Stories.docx"
TODAY = datetime.date(2026, 6, 18).strftime("%d %B %Y")

BRAND = RGBColor(0x4F, 0x46, 0xE5)      # indigo
INK   = RGBColor(0x1F, 0x29, 0x37)
MUTED = RGBColor(0x6B, 0x72, 0x80)
LABELBG = "EEF0F7"
HEADBG  = "4F46E5"

# ----------------------------------------------------------------------------
# USER STORY DATA  (grounded in the actual SafalEvents codebase)
# ----------------------------------------------------------------------------
EPICS = [
 {"id":"AUTH","name":"Authentication & Host Onboarding","desc":"Account creation, OTP verification, organization approval and sign-in across Guest, Host and Admin roles.","stories":[
   {"id":"US-AUTH-001","title":"Landing page & primary call-to-action","role":"first-time visitor","want":"land on a marketing page that explains SafalEvents and offers clear actions to explore events or become a host","benefit":"I understand the product and know how to start","pri":"Medium","deps":"—",
    "ac":["Landing page presents the value proposition, featured content and hero imagery.","Primary CTAs route to Explore Events and to Login/Sign-up.","Page is reachable at the site root (\"/\")."]},
   {"id":"US-AUTH-002","title":"Sign up as an individual host","role":"prospective host","want":"register with my first/last name, email, phone and password as an \"individual\" host","benefit":"I can create and manage my own events","pri":"High","deps":"US-AUTH-004",
    "ac":["Required fields are validated before a verification code is requested.","On submit, a signup session is created and a 6-digit OTP is sent to both email and SMS.","After successful OTP verification an individual host account is created with status ACTIVE and can sign in immediately."]},
   {"id":"US-AUTH-003","title":"Sign up as an organization host","role":"organization representative","want":"register my organization (name, type, website, country, city, state) and upload verification documents","benefit":"my organization can host events under a verified profile","pri":"High","deps":"US-AUTH-004, US-ADMIN-002",
    "ac":["Organization profile fields and at least one verification document are captured.","After OTP verification the account is created with status PENDING_ADMIN_APPROVAL.","The host cannot publish events until an admin approves the application.","Uploaded document names are persisted to the host record."]},
   {"id":"US-AUTH-004","title":"OTP verification (email + SMS)","role":"registering user","want":"confirm my identity with a one-time code sent to my email and phone","benefit":"only verified people can create accounts or RSVP","pri":"High","deps":"—",
    "ac":["A 6-digit code is generated with a 10-minute expiry.","Entering an expired code is rejected with a clear message and a resend option.","A maximum of 5 attempts is allowed; exceeding the limit invalidates the session.","Each attempt (success/failure) is written to the verification log with IP and device."]},
   {"id":"US-AUTH-005","title":"Sign in with email and password","role":"registered user (guest/host/admin)","want":"log in with my credentials and be routed to the correct dashboard for my role","benefit":"I get a role-appropriate experience","pri":"High","deps":"US-AUTH-002",
    "ac":["Valid credentials authenticate the user and set the current session.","Hosts land on the Host Dashboard, admins on the Admin Console, guests on the Guest Dashboard.","Invalid credentials are rejected with an error and no session is created."]},
   {"id":"US-AUTH-006","title":"Blocked sign-in for pending or suspended hosts","role":"organization host awaiting approval / suspended host","want":"be informed why I cannot access the host dashboard","benefit":"I understand my account state and next steps","pri":"High","deps":"US-AUTH-003, US-ADMIN-003",
    "ac":["A host in PENDING_ADMIN_APPROVAL is told the application is under review.","A suspended host is blocked and shown the suspension reason.","Approval or reinstatement restores normal access."]},
   {"id":"US-AUTH-007","title":"Log out","role":"signed-in user","want":"end my session securely","benefit":"my account is protected on shared devices","pri":"Medium","deps":"US-AUTH-005",
    "ac":["Logging out clears the current user session.","The user is redirected to the public landing page.","Protected dashboards are no longer accessible without signing in again."]},
 ]},

 {"id":"EVENT","name":"Event Creation & Management (Host)","desc":"Creating, configuring, publishing, editing, cloning and deleting events along with all RSVP and engagement settings.","stories":[
   {"id":"US-EVENT-001","title":"Create an event with core details","role":"host","want":"enter title, date, time, venue/location, city, state, description, cover image, theme and event type","benefit":"my event has a complete public listing","pri":"High","deps":"US-AUTH-005",
    "ac":["All core fields can be entered and a cover image and visual theme selected.","Saving creates an event and returns it to the host dashboard.","The event is assigned a unique ID and appears in Explore when published."]},
   {"id":"US-EVENT-002","title":"Set capacity and approval requirement","role":"host","want":"define maximum capacity and whether RSVPs require my approval","benefit":"I control attendance volume and who gets in","pri":"High","deps":"US-EVENT-001",
    "ac":["Capacity is stored and enforced during RSVP (over-capacity goes to waitlist).","When \"approval required\" is on, new RSVPs await host approval before becoming \"going\".","Defaults apply when fields are left blank."]},
   {"id":"US-EVENT-003","title":"Add custom RSVP questions","role":"host","want":"add custom questions guests must answer when they RSVP","benefit":"I collect information I need (allergies, song requests, startup details)","pri":"Medium","deps":"US-EVENT-001",
    "ac":["The host can add, edit and remove an ordered list of questions.","Questions appear in the RSVP flow and answers are stored per guest.","Answers are visible to the host in the guest list."]},
   {"id":"US-EVENT-004","title":"Configure privacy and guest-list visibility","role":"host","want":"set the event privacy (public/private/unlisted), whether the guest list is shown and how RSVP counts are displayed","benefit":"I control how much is public","pri":"Medium","deps":"US-EVENT-001",
    "ac":["Privacy setting controls discoverability in Explore.","Guest-list visibility toggle hides or shows attendees on the event page.","RSVP count display can be detailed or hidden."]},
   {"id":"US-EVENT-005","title":"Configure RSVP rules","role":"host","want":"set an RSVP deadline, max guests per RSVP, self-edit, self-cancellation, cancellation cut-off and whether a cancellation reason is required","benefit":"RSVP behaviour matches my event policy","pri":"High","deps":"US-EVENT-001",
    "ac":["RSVP deadline closes new RSVPs after the set date/time.","Max guests per RSVP caps the guest count a single registrant can bring.","Self-edit / self-cancellation toggles enable or disable guest self-service.","Cancellation cut-off (hours) and required-reason rule are enforced in the guest flow."]},
   {"id":"US-EVENT-006","title":"Configure engagement and confirmation settings","role":"host","want":"toggle comments, photo uploads and guest confirmation for the event","benefit":"I enable the right interaction level","pri":"Medium","deps":"US-EVENT-001",
    "ac":["Comments toggle shows/hides the comment board on the event page.","Photo-upload toggle enables/disables guest photo contributions.","Guest-confirmation toggle controls whether confirmation messaging is sent."]},
   {"id":"US-EVENT-007","title":"Enable paid tickets","role":"host","want":"turn on payments and set a ticket price for the event","benefit":"I can monetize attendance","pri":"Medium","deps":"US-EVENT-001, US-PAY-004",
    "ac":["Enabling payments reveals a ticket-price field.","Price is charged per guest during RSVP and an invoice is issued.","Disabling payments reverts the event to free RSVP."]},
   {"id":"US-EVENT-008","title":"Publish or save as draft","role":"host","want":"publish an event live or keep it as a draft","benefit":"I can prepare events before making them public","pri":"High","deps":"US-EVENT-001",
    "ac":["Draft events are not discoverable in Explore.","Published events appear publicly and accept RSVPs.","Status is clearly shown on the host dashboard."]},
   {"id":"US-EVENT-009","title":"Edit an event and notify guests","role":"host","want":"update event details after publishing and have guests informed of meaningful changes","benefit":"attendees always have accurate information","pri":"High","deps":"US-EVENT-008",
    "ac":["Edits are saved and reflected on the public event page.","A \"details updated\" notification can be dispatched to confirmed guests.","Increasing capacity automatically promotes waitlisted guests (FIFO).","The change is recorded in the audit trail."]},
   {"id":"US-EVENT-010","title":"Duplicate an event","role":"host","want":"clone an existing event as a starting point","benefit":"I can launch recurring/similar events quickly","pri":"Low","deps":"US-EVENT-001",
    "ac":["Cloning copies all settings into a new event titled \"… (Clone)\".","The clone is created as a Draft with its own ID and zeroed analytics.","RSVPs from the source are not copied."]},
   {"id":"US-EVENT-011","title":"Delete or cancel an event","role":"host","want":"remove an event and notify registered guests of cancellation","benefit":"I can call off events cleanly","pri":"Medium","deps":"US-EVENT-008",
    "ac":["Deleting an event removes its RSVPs, polls and comments.","A cancellation notification can be sent to confirmed guests.","The action is recorded in the audit trail."]},
   {"id":"US-EVENT-012","title":"Public host profile page","role":"guest","want":"view a host's public profile and their events","benefit":"I can judge credibility before RSVPing","pri":"Low","deps":"US-EVENT-001",
    "ac":["A host profile page is reachable at /host/{hostName}.","It shows host identity and their published events.","Links from event pages open the corresponding host profile."]},
 ]},

 {"id":"DISC","name":"Event Discovery (Guest)","desc":"Browsing, searching, filtering and viewing published events.","stories":[
   {"id":"US-DISC-001","title":"Browse published events","role":"visitor or guest","want":"see a gallery of published events","benefit":"I can find something to attend","pri":"High","deps":"US-EVENT-008",
    "ac":["Only published, discoverable events are listed.","Each card shows cover, title, date, location and key meta (rating, distance).","Selecting a card opens the event detail page."]},
   {"id":"US-DISC-002","title":"Search events by keyword","role":"guest","want":"search events by title or keyword","benefit":"I quickly find a specific event","pri":"Medium","deps":"US-DISC-001",
    "ac":["A search box filters the list as I type.","Matches consider title and relevant text fields.","An empty-state message shows when nothing matches."]},
   {"id":"US-DISC-003","title":"Filter and sort events","role":"guest","want":"filter by city/state/type and sort by distance, date or rating","benefit":"I see the most relevant events first","pri":"Medium","deps":"US-DISC-001",
    "ac":["City/state/type filters narrow the result set.","Sort options reorder results (distance, date, rating).","Filters and sort can be combined and cleared."]},
   {"id":"US-DISC-004","title":"View event detail page","role":"guest","want":"open an event to see full details, host, ratings and (if allowed) the guest list","benefit":"I have everything I need to decide","pri":"High","deps":"US-DISC-001",
    "ac":["Detail page at /e/{eventId} shows description, schedule, venue, cover and theme.","Rating, reviews count and host attribution are shown.","Guest list and RSVP counts honour the host's visibility settings.","A clear RSVP entry point is present."]},
   {"id":"US-DISC-005","title":"Count event views","role":"host (beneficiary)","want":"each event-page visit to be counted","benefit":"I can measure interest via analytics","pri":"Low","deps":"US-DISC-004",
    "ac":["Opening an event increments its view counter.","View totals feed the host analytics views metric."]},
 ]},

 {"id":"RSVP","name":"RSVP & Ticketing (Guest)","desc":"The end-to-end guest RSVP journey including verification, capacity, approval, payments and confirmations.","stories":[
   {"id":"US-RSVP-001","title":"Start an RSVP and choose a status","role":"guest","want":"begin an RSVP and select going, maybe or decline","benefit":"I can register my intent to attend","pri":"High","deps":"US-DISC-004",
    "ac":["RSVP flow is reachable at /rsvp/{eventId}.","Status options going/maybe/declined are selectable.","Name, email and phone are captured and validated."]},
   {"id":"US-RSVP-002","title":"Verify RSVP with OTP","role":"guest","want":"confirm my email/phone with a one-time code before my RSVP is recorded","benefit":"RSVPs are trustworthy and reduce spam","pri":"High","deps":"US-RSVP-001, US-AUTH-004",
    "ac":["A 6-digit code valid for 10 minutes is sent via email and SMS.","Up to 5 attempts are allowed; expiry/limits are enforced.","Previously verified guests for the same event may skip re-verification.","Attempts are logged."]},
   {"id":"US-RSVP-003","title":"Bring additional guests","role":"guest","want":"specify how many guests I am bringing up to the allowed maximum","benefit":"my whole party is counted","pri":"Medium","deps":"US-RSVP-001, US-EVENT-005",
    "ac":["Guest count cannot exceed the event's max-guests-per-RSVP.","Guest count is included in capacity calculations.","For paid events the charge multiplies by guest count."]},
   {"id":"US-RSVP-004","title":"Answer custom questions","role":"guest","want":"answer the host's custom questions during RSVP","benefit":"the host gets the details they requested","pri":"Medium","deps":"US-RSVP-001, US-EVENT-003",
    "ac":["All configured questions are presented in the flow.","Answers are stored against my RSVP record.","Answers are visible to the host."]},
   {"id":"US-RSVP-005","title":"Choose communication preferences","role":"guest","want":"select my preferred channel (email/SMS), reminder schedule and opt out of SMS","benefit":"I receive notifications the way I prefer","pri":"Low","deps":"US-RSVP-001",
    "ac":["Preferred channel and reminder schedule are saved on the RSVP.","SMS opt-out suppresses SMS notifications.","Preferences drive which confirmation/reminder messages are sent."]},
   {"id":"US-RSVP-006","title":"Automatic waitlisting when full","role":"guest","want":"be placed on a waitlist when the event is at capacity","benefit":"I still have a chance if a spot frees up","pri":"High","deps":"US-RSVP-003, US-EVENT-002",
    "ac":["If going-count + my guests exceeds capacity, my status becomes \"waitlist\".","I receive a waitlist notification with a booking ID.","If a spot opens I am promoted FIFO and notified automatically."]},
   {"id":"US-RSVP-007","title":"Approval-required registration","role":"guest","want":"submit my RSVP for host approval when the event requires it","benefit":"I understand my attendance is pending","pri":"Medium","deps":"US-RSVP-001, US-EVENT-002",
    "ac":["For approval-required events the RSVP is recorded as pending until the host decides.","I am notified when the host approves (going) or declines.","Approved RSVPs count toward capacity."]},
   {"id":"US-RSVP-008","title":"Pay for a ticket","role":"guest","want":"pay the ticket price (per guest) when RSVPing to a paid event","benefit":"I secure my paid spot","pri":"Medium","deps":"US-RSVP-003, US-EVENT-007",
    "ac":["Total charge equals ticket price × guest count.","A transaction record is created on successful payment.","An invoice/receipt with booking ID is sent to the guest."]},
   {"id":"US-RSVP-009","title":"Receive RSVP confirmation","role":"guest","want":"get a confirmation with my booking ID and a manage-RSVP link","benefit":"I have proof and can manage my RSVP","pri":"High","deps":"US-RSVP-002",
    "ac":["A confirmation is sent on the guest's preferred channel(s).","It includes event details, booking ID and a manage link.","Waitlisted guests receive the waitlist variant instead."]},
   {"id":"US-RSVP-010","title":"Enforce the RSVP deadline","role":"guest","want":"be prevented from RSVPing after the deadline","benefit":"expectations about closing are clear","pri":"Medium","deps":"US-EVENT-005",
    "ac":["After the RSVP deadline the RSVP entry point is closed.","A clear message explains that RSVPs are closed.","Already-registered guests can still manage their RSVP per host rules."]},
 ]},

 {"id":"GUEST","name":"Guest Dashboard & Self-Service","desc":"How guests view, manage and engage with their own RSVPs after registering.","stories":[
   {"id":"US-GUEST-001","title":"View my RSVPs","role":"guest","want":"see my upcoming and past events in one place","benefit":"I can keep track of what I'm attending","pri":"High","deps":"US-RSVP-009",
    "ac":["The dashboard lists the guest's RSVPs grouped by upcoming/past.","Each entry shows event, date, status and guest count.","Selecting an entry opens its details/management."]},
   {"id":"US-GUEST-002","title":"Access my digital pass / QR","role":"guest","want":"view my booking ID and QR pass for an event","benefit":"I can be checked in quickly on arrival","pri":"High","deps":"US-GUEST-001",
    "ac":["A QR pass / booking ID is available for confirmed (going) RSVPs.","The pass reflects the current RSVP status.","The pass is used by the host check-in scanner."]},
   {"id":"US-GUEST-003","title":"Edit my RSVP","role":"guest","want":"change my guest count or status when self-edit is allowed","benefit":"I can keep my RSVP accurate","pri":"Medium","deps":"US-GUEST-001, US-EVENT-005",
    "ac":["Editing is only available when the host enabled self-edit.","Changes respect max-guests and capacity (may trigger waitlist).","Updates are saved and reflected to the host."]},
   {"id":"US-GUEST-004","title":"Cancel my RSVP","role":"guest","want":"cancel my attendance subject to the host's rules","benefit":"I free my spot when I can't attend","pri":"Medium","deps":"US-GUEST-001, US-EVENT-005",
    "ac":["Cancellation is allowed only when enabled and before the cut-off window.","A reason is required when the host mandates it.","Cancelling frees capacity and can trigger waitlist promotion.","The host is notified of the cancellation."]},
   {"id":"US-GUEST-005","title":"Message the host","role":"guest","want":"send messages to the host of an event I'm attending and see replies","benefit":"I can ask questions and get answers","pri":"Medium","deps":"US-GUEST-001",
    "ac":["A two-way conversation thread exists per event/guest.","Sending a message flags it unread for the host and raises a host notification.","Host replies appear in my thread and unread indicators clear when read."]},
   {"id":"US-GUEST-006","title":"Receive host broadcasts and updates","role":"guest","want":"receive announcements, reminders and update notices from the host","benefit":"I stay informed about the event","pri":"Medium","deps":"US-RSVP-009, US-NOTIF-004",
    "ac":["Broadcasts and update/cancellation notices reach confirmed guests.","Messages honour my channel preferences.","Content is rendered from the relevant templates with event details."]},
   {"id":"US-GUEST-007","title":"Submit post-event feedback","role":"guest","want":"rate the event and leave comments after it ends","benefit":"I help the host improve and inform other guests","pri":"Medium","deps":"US-NOTIF-003",
    "ac":["A feedback form at /feedback/{eventId} captures a rating and comments.","Submission is stored and surfaced in host analytics.","Feedback requests are triggered after the event per schedule."]},
 ]},

 {"id":"GM","name":"Guest Management & Check-in (Host)","desc":"Host tools to manage attendees, approvals, waitlist, check-in, invitations and broadcasts.","stories":[
   {"id":"US-GM-001","title":"View and filter the RSVP list","role":"host","want":"see all RSVPs with status, guest counts and answers, with search/filter","benefit":"I understand and manage my attendees","pri":"High","deps":"US-EVENT-008",
    "ac":["RSVPs are listed with status (going/maybe/declined/waitlist) and details.","Custom-question answers are viewable per guest.","The list can be searched and filtered by status."]},
   {"id":"US-GM-002","title":"Approve or decline RSVPs","role":"host","want":"approve or decline pending RSVPs for approval-required events","benefit":"I curate my attendee list","pri":"High","deps":"US-EVENT-002, US-RSVP-007",
    "ac":["Pending RSVPs can be approved (→ going) or declined.","Approval respects remaining capacity.","The guest is notified of the decision and the action is audited."]},
   {"id":"US-GM-003","title":"Change RSVP status and auto-promote waitlist","role":"host","want":"manually change a guest's status and have the waitlist promote automatically when spots free up","benefit":"my attendee list stays correct and full","pri":"High","deps":"US-GM-001",
    "ac":["The host can change any RSVP's status.","Moving a \"going\" guest off the list frees capacity.","Freed capacity promotes waitlisted guests FIFO with notifications.","Each change is recorded in the audit trail."]},
   {"id":"US-GM-004","title":"Check guests in","role":"host","want":"check guests in via QR scan or manual toggle at the event","benefit":"I track real attendance","pri":"High","deps":"US-GUEST-002",
    "ac":["Scanning a guest's QR / toggling check-in marks them checked-in.","A host notification and a check-in confirmation are generated.","Check-in status is visible in the guest list and analytics.","The action is audited."]},
   {"id":"US-GM-005","title":"Add a manual invitation","role":"host","want":"add a guest directly to the going list without them self-registering","benefit":"I can include VIPs and offline invitees","pri":"Medium","deps":"US-GM-001",
    "ac":["The host enters guest name/email/phone and guest count.","The guest is added as \"going\" and flagged as a manual invite.","The invite counts toward capacity."]},
   {"id":"US-GM-006","title":"Broadcast to all guests","role":"host","want":"send an announcement to all guests of an event","benefit":"I can share urgent or important updates at once","pri":"Medium","deps":"US-GM-001",
    "ac":["A broadcast composes a message from a template and recipient set.","Delivery honours each guest's channel preferences.","The broadcast is recorded in the notification logs."]},
 ]},

 {"id":"ENG","name":"Engagement (Comments, Polls, Photos)","desc":"Interactive features on the event page for guests and host moderation.","stories":[
   {"id":"US-ENG-001","title":"Post and read comments","role":"guest","want":"post comments on an event and read others' comments","benefit":"I can engage with the community around the event","pri":"Medium","deps":"US-EVENT-006",
    "ac":["When comments are enabled, guests can post to the event board.","Comments display author, text and timestamp.","Comments are hidden when the host disables the feature."]},
   {"id":"US-ENG-002","title":"React to comments","role":"guest","want":"add emoji reactions to comments","benefit":"I can respond lightweight without posting","pri":"Low","deps":"US-ENG-001",
    "ac":["Guests can add emoji reactions to a comment.","Reaction counts update and persist.","Multiple reaction types are supported."]},
   {"id":"US-ENG-003","title":"Moderate comments","role":"host","want":"pin important comments and delete inappropriate ones","benefit":"I keep the board useful and clean","pri":"Medium","deps":"US-ENG-001",
    "ac":["The host can pin/unpin a comment so it surfaces at the top.","The host can delete any comment on their event.","Moderation actions take effect immediately."]},
   {"id":"US-ENG-004","title":"Create polls and vote","role":"host and guest","want":"the host to create polls and guests to vote on options","benefit":"the host gathers preferences and guests feel involved","pri":"Medium","deps":"US-EVENT-001",
    "ac":["The host can create a poll with a question and multiple options.","A guest can cast one vote and change it (vote moves between options).","Vote tallies update live and are visible on the event page."]},
   {"id":"US-ENG-005","title":"Guest photo uploads","role":"guest","want":"upload photos to an event when the host allows it","benefit":"attendees can share memories","pri":"Low","deps":"US-EVENT-006",
    "ac":["Photo upload is available only when enabled for the event.","Uploaded photos appear in the event gallery.","The feature is hidden when disabled."]},
 ]},

 {"id":"NOTIF","name":"Notifications & Templates Engine","desc":"Automated and scheduled messaging, per-event templates and the host activity inbox.","stories":[
   {"id":"US-NOTIF-001","title":"Automated RSVP confirmation","role":"guest","want":"receive an automatic confirmation when I RSVP","benefit":"I immediately know my registration succeeded","pri":"High","deps":"US-RSVP-009",
    "ac":["On RSVP, a confirmation is dispatched per event settings and guest channel.","Template variables (event, date, time, venue, booking ID, links) are rendered.","Every dispatch is recorded in the notification log with status."]},
   {"id":"US-NOTIF-002","title":"Scheduled pre-event reminders","role":"guest","want":"receive reminders before the event","benefit":"I don't forget to attend","pri":"High","deps":"US-NOTIF-001",
    "ac":["Reminder 1 fires at the configured offset (default 24h) before start.","An optional Reminder 2 fires at a second offset (default 3h) when enabled.","Reminders are sent once per guest and logged (no duplicates)."]},
   {"id":"US-NOTIF-003","title":"Post-event feedback request","role":"guest","want":"be invited to give feedback after the event","benefit":"I can share my experience","pri":"Medium","deps":"US-NOTIF-001, US-GUEST-007",
    "ac":["A feedback request is sent after the event by the configured delay.","Email and optional SMS variants are supported.","It is sent once per guest and logged."]},
   {"id":"US-NOTIF-004","title":"Update and cancellation notices","role":"guest","want":"be told when event details change or the event is cancelled","benefit":"I always have correct information","pri":"High","deps":"US-EVENT-009, US-EVENT-011",
    "ac":["Editing key details can dispatch a \"details updated\" notice.","Cancelling an event dispatches a cancellation notice to confirmed guests.","Both are rendered from templates and logged."]},
   {"id":"US-NOTIF-005","title":"Waitlist promotion notification","role":"guest","want":"be notified the moment I'm promoted from the waitlist","benefit":"I can claim my newly available spot","pri":"Medium","deps":"US-RSVP-006",
    "ac":["Promotion changes status to going and sends a promotion message.","The message includes pass/manage details.","The promotion is logged and audited."]},
   {"id":"US-NOTIF-006","title":"Host activity notification inbox","role":"host","want":"see a feed of new RSVPs, payments, check-ins, messages and capacity alerts","benefit":"I stay on top of event activity","pri":"High","deps":"US-EVENT-008",
    "ac":["Activity events create host notifications with type, title and message.","Unread notifications are indicated and can be marked read.","Notifications link to the relevant event."]},
   {"id":"US-NOTIF-007","title":"Customize per-event templates","role":"host","want":"edit the subject and body of my event's notification templates using variables","benefit":"messaging matches my voice and event","pri":"Medium","deps":"US-EVENT-001",
    "ac":["Per-event templates exist for confirmation, reminders, feedback, waitlist, updates, cancellation, invoice, etc.","Variables like {{event_name}}, {{guest_name}}, {{booking_id}} render correctly in a preview.","Saved templates are used for that event's dispatches."]},
   {"id":"US-NOTIF-008","title":"Notification outbox / delivery log","role":"host","want":"review a log of all messages sent for my event","benefit":"I can audit and troubleshoot communications","pri":"Medium","deps":"US-NOTIF-001",
    "ac":["The log lists messages with channel (email/SMS), type, subject, status and time.","Logs are filterable by event.","Both automated and manual dispatches appear."]},
 ]},

 {"id":"PAY","name":"Payments & Payouts (Host)","desc":"Paid ticketing, transactions, receipts and host bank/payout setup.","stories":[
   {"id":"US-PAY-001","title":"Configure paid ticketing","role":"host","want":"enable payments and set a ticket price","benefit":"I can sell tickets to my event","pri":"Medium","deps":"US-EVENT-007",
    "ac":["Payments can be toggled per event and a price set.","Pricing is reflected on the event page and RSVP flow.","Free events skip all payment steps."]},
   {"id":"US-PAY-002","title":"Record ticket transactions","role":"host","want":"have each paid RSVP recorded as a transaction","benefit":"I have an accurate sales record","pri":"Medium","deps":"US-PAY-001, US-RSVP-008",
    "ac":["A transaction captures guest, amount, guest count, total and method.","Transactions are listed per event.","Totals roll up into revenue analytics."]},
   {"id":"US-PAY-003","title":"View revenue and transactions","role":"host","want":"see total revenue and a transaction breakdown for an event","benefit":"I understand financial performance","pri":"Low","deps":"US-PAY-002",
    "ac":["A revenue summary aggregates completed transactions.","The transaction list is viewable and exportable.","Figures reconcile with the number of paid RSVPs."]},
   {"id":"US-PAY-004","title":"Add bank account for payouts","role":"host","want":"save my bank account details to receive payouts","benefit":"I can be paid for ticket sales","pri":"Low","deps":"US-AUTH-005",
    "ac":["Bank details are saved to my host profile.","Saved details can be viewed/updated.","Payout setup is required before monetizing (informational gate)."]},
   {"id":"US-PAY-005","title":"Issue ticket invoice / receipt","role":"guest","want":"receive a receipt for my ticket purchase","benefit":"I have proof of payment","pri":"Medium","deps":"US-RSVP-008",
    "ac":["A receipt with total paid and booking ID is sent on purchase.","The receipt links to my manage-RSVP/pass.","The dispatch is logged."]},
 ]},

 {"id":"ANALYTICS","name":"Analytics & Reporting (Host)","desc":"Insights for hosts across views, RSVPs, attendance, feedback and revenue.","stories":[
   {"id":"US-ANALYTICS-001","title":"Views and interest funnel","role":"host","want":"see how many people viewed my event versus how many RSVP'd","benefit":"I can gauge conversion and reach","pri":"Medium","deps":"US-DISC-005",
    "ac":["View totals are shown per event.","Views vs RSVPs are presented as a simple funnel/ratio.","Metrics update as activity occurs."]},
   {"id":"US-ANALYTICS-002","title":"RSVP and attendance breakdown","role":"host","want":"see counts by status and the check-in rate","benefit":"I can plan logistics and measure turnout","pri":"Medium","deps":"US-GM-001, US-GM-004",
    "ac":["Counts for going/maybe/declined/waitlist are displayed.","Check-in rate (checked-in vs going) is calculated.","Guest-count totals account for additional guests."]},
   {"id":"US-ANALYTICS-003","title":"Feedback and ratings summary","role":"host","want":"see aggregated ratings and feedback comments","benefit":"I can evaluate guest satisfaction","pri":"Low","deps":"US-GUEST-007",
    "ac":["Average rating and response count are shown.","Individual feedback comments are listed.","Summary reflects only submitted responses."]},
   {"id":"US-ANALYTICS-004","title":"Revenue summary","role":"host","want":"see total revenue for paid events","benefit":"I can measure commercial success","pri":"Low","deps":"US-PAY-002",
    "ac":["Revenue totals are derived from completed transactions.","Figures are shown only for payment-enabled events.","Totals match the transaction list."]},
 ]},

 {"id":"ADMIN","name":"Admin Console","desc":"Platform administration: host approvals, suspensions, system templates, settings, verification logs and audit.","stories":[
   {"id":"US-ADMIN-001","title":"View users and hosts","role":"admin","want":"see all registered users/hosts with their type and status","benefit":"I have oversight of the platform's accounts","pri":"High","deps":"US-AUTH-005",
    "ac":["All users are listed with role, host type and status.","The list can be searched/filtered.","Selecting a host shows their profile and documents."]},
   {"id":"US-ADMIN-002","title":"Approve or reject host applications","role":"admin","want":"review pending organization hosts and approve or reject with a reason","benefit":"only legitimate organizations can host","pri":"High","deps":"US-AUTH-003",
    "ac":["Pending applications and their uploaded documents are reviewable.","Approval sets status ACTIVE; rejection stores a reason.","The applicant is emailed the decision (approval/rejection with reason)."]},
   {"id":"US-ADMIN-003","title":"Suspend or reinstate hosts","role":"admin","want":"suspend a host (with reason) and later reinstate them","benefit":"I can enforce policy and protect guests","pri":"High","deps":"US-ADMIN-001",
    "ac":["Suspending records a reason and blocks the host's access.","Reinstating clears the suspension and restores access.","Suspension state is reflected at sign-in."]},
   {"id":"US-ADMIN-004","title":"Manage system-wide templates","role":"admin","want":"edit platform default notification templates with versioning","benefit":"consistent baseline messaging across all events","pri":"Medium","deps":"US-NOTIF-007",
    "ac":["System templates can be edited (subject/body).","Each save bumps the version and appends a change-log entry.","Version history is viewable."]},
   {"id":"US-ADMIN-005","title":"Configure platform settings","role":"admin","want":"set platform name, primary color, from-name, support email, SMS sender ID and RSVP defaults","benefit":"I control global branding and defaults","pri":"Medium","deps":"US-AUTH-005",
    "ac":["Settings persist and apply platform-wide.","Default RSVP deadline and self-cancellation defaults seed new events.","Branding values are reflected where used."]},
   {"id":"US-ADMIN-006","title":"Review verification logs","role":"admin","want":"audit OTP verification attempts with outcome, IP and device","benefit":"I can detect abuse and support users","pri":"Medium","deps":"US-AUTH-004, US-RSVP-002",
    "ac":["All signup and RSVP OTP attempts are logged.","Each entry shows email/phone, outcome, IP and device.","Logs are viewable and filterable."]},
   {"id":"US-ADMIN-007","title":"View the immutable audit trail","role":"admin","want":"see a chronological, tamper-evident record of key actions","benefit":"I have accountability and traceability","pri":"Medium","deps":"US-PLAT-001",
    "ac":["Key actions (create/update/RSVP/check-in/promotion/admin) are recorded.","Each entry has timestamp, actor and action.","Entries are append-only (not editable from the UI)."]},
   {"id":"US-ADMIN-008","title":"Oversee all events","role":"admin","want":"view all events and platform-level analytics","benefit":"I can monitor and moderate activity","pri":"Low","deps":"US-ADMIN-001",
    "ac":["All events across hosts are visible to the admin.","Platform-level metrics are summarized.","Problematic events can be acted upon (via host suspension/template control)."]},
 ]},

 {"id":"APPROVAL","name":"RSVP Approval Workflow","desc":"Organizer approval of RSVPs with an Under Approval -> Approved/Rejected lifecycle (orthogonal to the going/maybe/declined response), capacity-driven waitlisting that also awaits approval, lifecycle emails and a full activity history.","stories":[
   {"id":"US-APPROVAL-001","title":"RSVP enters Under Approval","role":"guest","want":"have my RSVP held as 'Under Approval' when the event requires organizer approval","benefit":"I know my spot is pending the host's decision","pri":"High","deps":"US-EVENT-002, US-RSVP-001",
    "ac":["When the event has 'approval required', a new RSVP is recorded with approvalState UNDER_APPROVAL (separate from the going/maybe/declined response).","The guest screen clearly shows an 'Under Approval' badge for that event.","The host (and staff with notification permission) are notified a new RSVP awaits review.","The submission is written to the event's activity history; a guest cannot self-approve.","Re-submitting/editing the RSVP returns it to UNDER_APPROVAL and logs the change."]},
   {"id":"US-APPROVAL-002","title":"Approve or reject a pending RSVP","role":"host or permitted staff","want":"review pending RSVPs and approve or reject each (optionally with a reason)","benefit":"I curate exactly who is admitted","pri":"High","deps":"US-APPROVAL-001, US-STAFF-002",
    "ac":["Approve sets approvalState APPROVED and admits the guest (status going); reject sets REJECTED with an optional reason.","Only the host, or staff whose role grants Approve RSVP, can decide; everyone else sees the queue read-only.","The guest is emailed the decision and the action (who/what/when) is written to history.","A bulk 'Approve all pending' action is available."]},
   {"id":"US-APPROVAL-003","title":"Re-open a rejected RSVP","role":"host","want":"move a rejected RSVP back into the pending queue","benefit":"I can reverse a decision if circumstances change","pri":"Low","deps":"US-APPROVAL-002",
    "ac":["A rejected RSVP can be re-opened to UNDER_APPROVAL.","The re-open is recorded in the activity history.","The guest's screen reflects the restored 'Under Approval' state."]},
   {"id":"US-APPROVAL-004","title":"Capacity-full RSVP waitlisted under approval","role":"guest","want":"be waitlisted and held for approval when the event is full","benefit":"I still have a chance if a seat frees up","pri":"High","deps":"US-EVENT-002, US-RSVP-006",
    "ac":["When capacity is reached, a new RSVP becomes 'waitlist' and is held UNDER_APPROVAL.","Waitlisted guests are not auto-admitted; the host approves them to allow them in.","When a confirmed guest cancels, the host is notified that a spot opened to review the waitlist.","Approving a waitlisted guest admits them (status going)."]},
   {"id":"US-APPROVAL-005","title":"Approval lifecycle emails","role":"guest and host","want":"the right email at each step of the approval lifecycle","benefit":"both sides always know the current state","pri":"High","deps":"US-NOTIF-001",
    "ac":["Guest receives 'received - pending approval' on submit (no confirmation/ticket yet).","Guest receives an 'approved' email on approval and a 'not approved' email (with reason) on rejection.","Host receives a 'new RSVP awaiting approval' notice with a link to the Manage -> Guests page.","Every message is recorded in the outbox / history."]},
   {"id":"US-APPROVAL-006","title":"Approval status across guest screens","role":"guest","want":"see my approval state consistently wherever my RSVP appears","benefit":"there is no ambiguity about whether I'm in","pri":"Medium","deps":"US-GUEST-001",
    "ac":["Under Approval / Waitlisted / Approved / Not Approved render on the event page, the ticket and the dashboard.","The QR pass / 'View pass' is available only once approved.","Rejected RSVPs show the reason when one was provided."]},
   {"id":"US-APPROVAL-007","title":"RSVP & activity history","role":"host or permitted staff","want":"a complete chronological record of all RSVP activity for an event","benefit":"I have a full, auditable trail","pri":"Medium","deps":"US-PLAT-001",
    "ac":["History shows submissions/edits, status changes (who/when) and emails sent.","Entries are append-only (immutable).","History is visible to the host and to staff whose role grants History access."]},
 ]},

 {"id":"STAFF","name":"Staff, Roles & Permissions","desc":"Bringing team members onto an event with controlled access via named, module/action-level roles, enforced in the UI and on every request (default-deny).","stories":[
   {"id":"US-STAFF-001","title":"Invite a staff member with a role","role":"host (or staff with Manage Staff permission)","want":"invite a team member by email and assign them a role for my event","benefit":"others can help run the event with the right access","pri":"High","deps":"US-EVENT-008",
    "ac":["The host enters name + email and selects a role; an invite email is sent.","A shareable, unique Invite ID is generated and shown for that staff member.","The staff member appears in the event's staff list as INVITED until accepted.","The invite/assignment is recorded in history."]},
   {"id":"US-STAFF-002","title":"Define roles and permissions (Add Role)","role":"host","want":"create/edit named roles by toggling permissions for every screen and action","benefit":"I grant exactly the access each helper needs","pri":"High","deps":"US-EVENT-008",
    "ac":["An Add/Edit Role screen lists every module/action: view guests, approve/reject RSVPs, edit guests, export, gate check-in, view/reply messaging, view history, view/edit settings, manage staff.","Default-deny: anything not granted is hidden in the UI and blocked server-side.","Built-in roles are provided (Coordinator, Front-desk, QR Scanner, Viewer) and custom roles can be created.","Editing a role updates all staff assigned to it; built-in roles cannot be deleted."]},
   {"id":"US-STAFF-003","title":"Role-limited experience for staff","role":"staff member","want":"see only the tabs and actions my role permits","benefit":"I can do exactly my job and nothing more","pri":"High","deps":"US-STAFF-002, US-ACCESS-003",
    "ac":["The dashboard and event pages render only the screens/actions the role grants.","Attempting a non-permitted feature (e.g. via direct URL) is denied.","The host is implicitly full-access (owner) on their own events."]},
   {"id":"US-STAFF-004","title":"Accept invite and operate within limits","role":"staff member","want":"accept my invite and access only the event(s) I was invited to","benefit":"I get straight to my assigned work","pri":"Medium","deps":"US-STAFF-001",
    "ac":["Accepting the invite activates the assignment (INVITED -> ACTIVE).","The staff member is scoped to only their assigned event(s).","Their permitted screens follow their role (US-STAFF-003)."]},
   {"id":"US-STAFF-005","title":"Manage and revoke staff","role":"host","want":"remove a staff member or change their role","benefit":"I keep the team and its access current","pri":"Medium","deps":"US-STAFF-001",
    "ac":["A staff member can be removed/revoked (their Invite ID stops working).","Changing a role re-applies to that member's access immediately.","Add/remove/role-change actions are audited."]},
 ]},

 {"id":"CHECKIN","name":"Gate Check-in & QR Scanner","desc":"Verifying and marking attendance at the venue gate by scanning each guest's QR pass, with a dedicated QR Scanner role.","stories":[
   {"id":"US-CHECKIN-001","title":"QR Scanner role","role":"host","want":"assign a role whose only screen is gate Check-in","benefit":"door staff can scan without seeing guest editing, messaging or settings","pri":"Medium","deps":"US-STAFF-002",
    "ac":["A built-in 'QR Scanner' role exposes only the Check-in module.","Check-in is a role permission that can be combined with others if the host wishes.","A QR Scanner signs in (US-ACCESS-003) and lands directly on Check-in."]},
   {"id":"US-CHECKIN-002","title":"Scan a pass and verify against the approved list","role":"gate staff","want":"scan a guest's QR / pass and have it validated against the event's approved guests","benefit":"only valid, approved guests are admitted","pri":"High","deps":"US-CHECKIN-001, US-APPROVAL-002",
    "ac":["Valid and not yet checked in -> mark Arrived and show guest name + RSVP details (e.g. number of attendees).","Already checked in -> warn 'Already scanned at [time]'.","Invalid / not approved / still under approval / wrong event -> 'Not valid for this event' (entry denied).","Passes are single-event; a pass for another event is rejected."]},
   {"id":"US-CHECKIN-003","title":"Live arrivals, emails and logging","role":"host (beneficiary)","want":"each scan to update counts, notify and be logged","benefit":"I see real-time turnout and have an audit trail","pri":"Medium","deps":"US-CHECKIN-002, US-NOTIF-006",
    "ac":["A successful check-in updates the live arrival count on the host dashboard.","Each scan (guest, scanner, timestamp) is written to the activity history.","A welcome/check-in confirmation is sent to the guest."]},
 ]},

 {"id":"MSG","name":"Per-Event Messaging","desc":"Host-configurable in-app guest <-> host messaging that appears to guests only when the host enables it for that event.","stories":[
   {"id":"US-MSG-001","title":"Configure messaging per event","role":"host","want":"turn guest messaging on or off for a specific event","benefit":"I control whether guests can message me for each event","pri":"Medium","deps":"US-EVENT-001",
    "ac":["An 'Allow guest messaging' toggle is saved on the event (per-event, independent of other events).","Turning it off hides messaging from guests; existing thread history is preserved.","The default for new events is configurable."]},
   {"id":"US-MSG-002","title":"Guest <-> host messaging when enabled","role":"guest","want":"message the host in-app for an event that allows it and see replies","benefit":"I can ask questions and get answers","pri":"Medium","deps":"US-MSG-001, US-GUEST-005",
    "ac":["The 'Message host' option appears to the guest only when the host enabled messaging for that event.","Messages are delivered in-app as a per-event thread; the host (and permitted staff) can reply.","When messaging is off, the guest sees no messaging option for that event."]},
   {"id":"US-MSG-003","title":"Staff messaging by permission","role":"staff member","want":"read and reply to guest messages only if my role allows it","benefit":"messaging access matches my responsibilities","pri":"Low","deps":"US-STAFF-002",
    "ac":["Viewing messages requires the messaging-view permission; replying requires messaging-reply.","Staff without messaging permission do not see the Messaging screen.","Replies are attributed and appear in the guest's thread."]},
 ]},

 {"id":"ACCESS","name":"Account Types: Signup & Login","desc":"Choosing Guest vs Host at signup, a unified email/phone login that routes by account type, and a separate staff Invite-ID login.","stories":[
   {"id":"US-ACCESS-001","title":"Sign up as Guest or Host","role":"new user","want":"clearly choose to register as a Guest or as a Host","benefit":"I land in the right experience for what I want to do","pri":"High","deps":"US-AUTH-004",
    "ac":["The signup screen offers two clear options: register as Guest or as Host (email or phone + OTP).","The account is provisioned with the chosen type: Host -> Host flow; Guest -> Guest flow.","Staff do not sign up here - they join via an Invite ID (US-ACCESS-003).","A single person may be a host for their events and a guest for others; account type sets the default flow."]},
   {"id":"US-ACCESS-002","title":"Unified login routes by account type","role":"registered Host or Guest","want":"log in with email or phone and be taken to the correct flow","benefit":"one simple login does the right thing","pri":"High","deps":"US-AUTH-005",
    "ac":["Email/phone authentication (OTP/link or password) signs the user in.","Hosts are routed to the Host flow and Guests to the Guest flow based on account type.","Invalid credentials are rejected without creating a session."]},
   {"id":"US-ACCESS-003","title":"Login as Staff via Invite ID","role":"invited staff member","want":"a 'Login as Staff' path where I enter my Invite ID and email/phone","benefit":"I reach my assigned event with only my role's screens","pri":"High","deps":"US-STAFF-001, US-STAFF-003",
    "ac":["A 'Login as Staff' button reveals Invite ID + email/phone fields.","The Invite ID is validated against the email/phone it was issued to.","A wrong/expired/revoked Invite ID is rejected with a clear message and a resend option.","On success the staff lands in their invited event(s) with only role-permitted screens."]},
 ]},

 {"id":"MOBILE","name":"Mobile App (Browse-before-login & Parity)","desc":"The SafalEvents mobile app (Expo / React Native): browse as a guest before logging in, gated-action auth that returns to intent, full host/guest/staff parity, and a phone-frame web build.","stories":[
   {"id":"US-MOBILE-001","title":"Browse before login (Guest Mode)","role":"newly installed user","want":"explore the app without logging in or registering","benefit":"I get a feel for it before committing to an account","pri":"High","deps":"—",
    "ac":["After the splash the app opens into a browse/Guest Mode home - no login wall.","The user can scroll events and open event detail screens without an account.","Guest Mode is read-only - nothing is written until the user authenticates."]},
   {"id":"US-MOBILE-002","title":"Gated actions trigger auth","role":"guest-mode user","want":"be sent to Login/Register only when I tap something that needs an account","benefit":"I'm never blocked until it's actually necessary","pri":"High","deps":"US-MOBILE-001, US-ACCESS-001",
    "ac":["RSVP, Message host, Save/favorite, Create event and My Dashboard route to the Login/Register screen.","A persistent 'Login / Sign up' entry stays visible while browsing.","The auth screen offers Register/Login as Guest or Host and a 'Login as Staff' button."]},
   {"id":"US-MOBILE-003","title":"Return to intent (and cancel)","role":"guest-mode user","want":"be brought back to the exact action I tapped after I authenticate","benefit":"the flow feels seamless and I don't lose my place","pri":"High","deps":"US-MOBILE-002",
    "ac":["After successful login/registration the app returns to the screen/action that triggered auth (e.g. the RSVP reopens).","Dismissing the auth screen (back/swipe) returns to Guest Mode with nothing changed.","Browsing context (selected event) is preserved across the auth flow."]},
   {"id":"US-MOBILE-004","title":"Correct flow after auth","role":"authenticated user","want":"be routed to my role's experience after signing in","benefit":"I see what's relevant to me","pri":"High","deps":"US-ACCESS-002, US-STAFF-003",
    "ac":["Host -> Host flow; Guest -> Guest flow.","Staff -> their invited event(s) with only role-permitted screens.","A QR Scanner lands directly on Check-in."]},
   {"id":"US-MOBILE-005","title":"Persistent entry, remembered session & deep links","role":"returning or linked user","want":"a discoverable login entry, a remembered session, and deep links that respect gating","benefit":"returning and shared-link journeys are smooth","pri":"Medium","deps":"US-MOBILE-001",
    "ac":["A Login/Sign up entry is always visible during browsing.","The device remembers the session after first login so returning users skip Guest Mode.","Deep links / push taps open the target screen in Guest Mode; if it needs auth, the same gated redirect applies, then returns to that screen."]},
   {"id":"US-MOBILE-006","title":"Host/Guest/Staff parity on mobile","role":"any user","want":"the mobile app to mirror the web feature set","benefit":"I can run everything from my phone","pri":"Medium","deps":"US-CHECKIN-002",
    "ac":["Host: dashboard, events, RSVP approvals + waitlist, staff & roles, per-event settings.","Guest: tickets with approval/waitlist states, explore, RSVP, ticket pass, messaging when enabled.","Staff: a QR-scan check-in that marks guests arrived and updates the host dashboard live."]},
   {"id":"US-MOBILE-007","title":"Phone-frame web build","role":"reviewer","want":"open the mobile app in a browser and have it look like a phone","benefit":"I can preview the mobile UI without a device","pri":"Low","deps":"US-PLAT-004",
    "ac":["The app runs on the web via react-native-web inside a phone-frame device mockup.","It is deployable as a static web build to a host such as Netlify.","Layout and theming match the handset experience."]},
 ]},

 {"id":"PLAT","name":"Platform & Cross-Cutting","desc":"Non-functional and cross-cutting capabilities underpinning the product.","stories":[
   {"id":"US-PLAT-001","title":"Immutable audit logging","role":"platform owner","want":"critical actions to be automatically written to an append-only audit trail","benefit":"the platform is accountable and traceable","pri":"Medium","deps":"—",
    "ac":["Create/update/delete, RSVP changes, check-ins, promotions and admin actions are logged.","Logs capture actor, action and timestamp.","Logs are surfaced to admins (US-ADMIN-007)."]},
   {"id":"US-PLAT-002","title":"OTP security controls","role":"platform owner","want":"one-time codes to expire, limit attempts and be fully logged","benefit":"verification is secure against brute force and reuse","pri":"High","deps":"US-AUTH-004, US-RSVP-002",
    "ac":["Codes expire after 10 minutes.","A 5-attempt limit invalidates the session when exceeded.","All attempts are logged with metadata."]},
   {"id":"US-PLAT-003","title":"Role-based access control","role":"platform owner","want":"each role to see only the screens and actions appropriate to it","benefit":"users cannot access unauthorized functions","pri":"High","deps":"US-AUTH-005",
    "ac":["Guest, Host and Admin are routed to their respective dashboards.","Host/Admin tools are not exposed to guests.","Unauthenticated users are redirected to sign-in for protected areas."]},
   {"id":"US-PLAT-004","title":"Responsive, themed UI","role":"any user","want":"a responsive interface that works on desktop and mobile with event themes","benefit":"I have a good experience on any device","pri":"Medium","deps":"—",
    "ac":["Core screens adapt to mobile and desktop widths.","Event visual themes render consistently.","Interactive controls are usable on touch devices."]},
   {"id":"US-PLAT-005","title":"Reliable state persistence","role":"any user","want":"my data (events, RSVPs, settings) to persist across sessions","benefit":"I don't lose my work or registrations","pri":"High","deps":"—",
    "ac":["Created events, RSVPs and settings persist between visits.","Data migrations backfill new fields without data loss.","State is restored on return to the app."]},
 ]},
]

# ----------------------------------------------------------------------------
# DOCX BUILD HELPERS
# ----------------------------------------------------------------------------
def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), hexcolor)
    tcPr.append(shd)

def set_cell_text(cell, text, bold=False, color=None, size=None, white=False):
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    if size: run.font.size = Pt(size)
    if white: run.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
    elif color: run.font.color.rgb = color
    return p

def add_bullets(cell, items):
    cell.text = ''
    for i, it in enumerate(items):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        p.style = doc.styles['List Bullet']
        p.add_run(it)

doc = Document()

# base font
normal = doc.styles['Normal']
normal.font.name = 'Calibri'; normal.font.size = Pt(10.5)
normal.font.color.rgb = INK

for h, sz in [('Heading 1', 16), ('Heading 2', 13.5), ('Heading 3', 11.5), ('Title', 30)]:
    st = doc.styles[h]
    st.font.name = 'Calibri'
    if h != 'Title':
        st.font.color.rgb = BRAND

# ---------- TITLE PAGE ----------
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run('\n\n\n\nSafalEvents'); r.font.size = Pt(40); r.bold = True; r.font.color.rgb = BRAND
sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
rs = sub.add_run('End-to-End User Stories'); rs.font.size = Pt(20); rs.font.color.rgb = INK
sub2 = doc.add_paragraph(); sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
rs2 = sub2.add_run('Product Requirements — Functional User Stories with Acceptance Criteria'); rs2.font.size = Pt(11); rs2.font.color.rgb = MUTED

doc.add_paragraph('\n\n')
info = doc.add_table(rows=5, cols=2); info.alignment = WD_TABLE_ALIGNMENT.CENTER; info.style = 'Table Grid'
info.columns[0].width = Inches(1.8); info.columns[1].width = Inches(3.6)
meta = [('Product', 'SafalEvents — Event hosting, RSVP & ticketing platform'),
        ('Document', 'End-to-End User Stories'),
        ('Version', '2.0'),
        ('Date', TODAY),
        ('Status', 'Draft for review')]
for i,(k,v) in enumerate(meta):
    set_cell_text(info.rows[i].cells[0], k, bold=True); shade(info.rows[i].cells[0], LABELBG)
    set_cell_text(info.rows[i].cells[1], v)
doc.add_page_break()

# ---------- HOW TO READ ----------
doc.add_heading('How to read this document', level=1)
doc.add_paragraph('This document captures the complete, end-to-end functional scope of SafalEvents as user stories. '
                  'Stories are grouped into epics (modules). Each story has a unique ID, a role-based narrative '
                  '(As a … I want … so that …), testable acceptance criteria, a priority and dependencies on other stories.')
conv = doc.add_table(rows=1, cols=2); conv.style = 'Table Grid'
hdr = conv.rows[0].cells
set_cell_text(hdr[0], 'Convention', bold=True, white=True); shade(hdr[0], HEADBG)
set_cell_text(hdr[1], 'Meaning', bold=True, white=True); shade(hdr[1], HEADBG)
for k,v in [('Story ID', 'US-<EPIC>-<NNN>, e.g. US-RSVP-006. IDs are stable references for traceability.'),
            ('Priority', 'High = core / must-have, Medium = important, Low = nice-to-have.'),
            ('Dependencies', 'Other story IDs that should be delivered first; "—" means none.'),
            ('Roles', 'Guest, Host, Admin (plus Visitor and Platform Owner for cross-cutting items).')]:
    row = conv.add_row().cells
    set_cell_text(row[0], k, bold=True); shade(row[0], LABELBG)
    set_cell_text(row[1], v)

# ---------- EPIC SUMMARY ----------
doc.add_heading('Epic summary', level=1)
es = doc.add_table(rows=1, cols=4); es.style = 'Table Grid'
hc = es.rows[0].cells
for i,h in enumerate(['Epic ID','Epic','# Stories','Scope']):
    set_cell_text(hc[i], h, bold=True, white=True); shade(hc[i], HEADBG)
for ep in EPICS:
    c = es.add_row().cells
    set_cell_text(c[0], 'EPIC-'+ep['id'], bold=True)
    set_cell_text(c[1], ep['name'])
    set_cell_text(c[2], str(len(ep['stories'])))
    set_cell_text(c[3], ep['desc'])

# ---------- TRACEABILITY INDEX ----------
doc.add_heading('Traceability index (all stories)', level=1)
idx = doc.add_table(rows=1, cols=4); idx.style = 'Table Grid'
ic = idx.rows[0].cells
for i,h in enumerate(['Story ID','Title','Priority','Depends on']):
    set_cell_text(ic[i], h, bold=True, white=True); shade(ic[i], HEADBG)
total = 0
for ep in EPICS:
    for s in ep['stories']:
        total += 1
        c = idx.add_row().cells
        set_cell_text(c[0], s['id'], bold=True)
        set_cell_text(c[1], s['title'])
        set_cell_text(c[2], s['pri'])
        set_cell_text(c[3], s['deps'])
doc.add_page_break()

# ---------- DETAILED STORIES ----------
doc.add_heading('Detailed user stories', level=1)
for ep in EPICS:
    doc.add_heading('EPIC-%s — %s' % (ep['id'], ep['name']), level=2)
    doc.add_paragraph(ep['desc'])
    for s in ep['stories']:
        doc.add_heading('%s — %s' % (s['id'], s['title']), level=3)
        tbl = doc.add_table(rows=0, cols=2); tbl.style = 'Table Grid'
        tbl.columns[0].width = Inches(1.5); tbl.columns[1].width = Inches(5.0)
        def row(label, value=None, bullets=None):
            cells = tbl.add_row().cells
            set_cell_text(cells[0], label, bold=True); shade(cells[0], LABELBG)
            if bullets is not None:
                add_bullets(cells[1], bullets)
            else:
                set_cell_text(cells[1], value)
        row('Story ID', s['id'])
        row('Epic', 'EPIC-%s — %s' % (ep['id'], ep['name']))
        row('Priority', s['pri'])
        row('User story', 'As a %s, I want to %s, so that %s.' % (s['role'], s['want'], s['benefit']))
        row('Acceptance criteria', bullets=s['ac'])
        row('Dependencies', s['deps'])
        doc.add_paragraph('')

# ---------- FOOTER (page numbers) ----------
def add_page_number(footer_par):
    run = footer_par.add_run()
    fldStart = OxmlElement('w:fldChar'); fldStart.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'), 'preserve'); instr.text = 'PAGE'
    fldEnd = OxmlElement('w:fldChar'); fldEnd.set(qn('w:fldCharType'), 'end')
    run._r.append(fldStart); run._r.append(instr); run._r.append(fldEnd)

sec = doc.sections[0]
foot = sec.footer.paragraphs[0]; foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = foot.add_run('SafalEvents — End-to-End User Stories  ·  Page '); fr.font.size = Pt(8); fr.font.color.rgb = MUTED
add_page_number(foot)

doc.save(OUT)
print('Stories:', total)
print('Saved:', OUT)
